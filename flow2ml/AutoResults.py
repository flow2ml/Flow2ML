from sklearn.metrics import roc_auc_score, roc_curve
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from sklearn.metrics import plot_confusion_matrix
from sklearn.metrics import plot_precision_recall_curve
from docx import Document
from docx.shared import Inches
import os


class Auto_Results:
    def __init__(self,model,X_test,Y_test):
        '''
        Initialises the class variables.
        '''
        self.model = model          #stores the model
        self.x_test = X_test        #stores the X_test (test data) 
        self.prediction = self.model.predict_proba(X_test)     #stores the probabilities for roc
        self.y_test = Y_test    #stores the Y_test (labels)
        try:                        #try creating a Results directory and show excpetion if it already exists 
            p = os.getcwd()
            p = os.path.join(p,"Results")
            self.results_path = p
            os.mkdir(self.results_path)
        except:
            print("Results Folder already exists")

    def roc_curve(self,figure_name="roc.jpeg"):
        '''
        Plotting the Roc curve for the Input model
        '''
        i= self.prediction.shape
        try:
            if (i[1]>1):
                self.prediction = self.prediction[:, 1]
        except:
            print("Predictions vector doesn't need slicing")
        auc_lr = roc_auc_score(self.y_test, self.prediction)  
        fpr_lr, tpr_lr, thresholds_lr = roc_curve(self.y_test, self.prediction)
        plt.figure(figsize=(50, 20))
        plt.plot(fpr_lr, tpr_lr)
        plt.title('ROC Curve', fontsize=70)
        plt.xlabel('False Positive Rate', fontsize=65)
        plt.ylabel('True Positive Rate', fontsize=65)
        plt.savefig(self.results_path+'/'+figure_name)
    
    def confusion_matrix(self):
        '''
        Plotting the confusion matrix
        '''
        plt.figure(figsize=(50,20))
        plot_confusion_matrix(self.model,self.x_test,self.y_test)
        plt.title('Confusion Matrix', size = 35)
        plt.savefig(self.results_path+'/confusion_matrix.jpeg')
    
    def precision_recall_curve(self,figure_name="prc.jpeg"):
        '''
        Plotting the precision recall curve 
        '''
        plt.figure(figsize=(50,20))
        plot_precision_recall_curve(self.model,self.x_test,self.y_test)
        plt.title('Precision Recall Curve', size = 35)
        plt.savefig(self.results_path+'/'+figure_name)

    def get_results_pdf(self,file_name="report"):
        '''
        Saves all the plots with their default name and 
        creates a report.docx file in the Results Folder
        '''
        self.roc_curve()
        self.confusion_matrix()
        self.precision_recall_curve()
        doc = Document()
        doc.add_heading('Model Report')
        doc.add_heading('ROC Curve',level = 2)
        doc.add_paragraph('''
An ROC curve (receiver operating characteristic curve) is a graph showing the performance of a classification model at all classification thresholds.
This curve plots two parameters:
        True Positive Rate
        False Positive Rate
An ROC curve plots TPR vs. FPR at different classification thresholds. Lowering the classification threshold classifies more items as positive, thus increasing both False Positives and True Positives.
The ROC curve for the Input Model is As Follows.
''')
        doc.add_picture(self.results_path+'/roc.jpeg',width=Inches(6.0))
        doc.add_paragraph('''  ''')
        doc.add_heading('Confusion Matrix', level=2)
        doc.add_paragraph('''
A confusion matrix is a table that is often used to describe the performance of a classification model (or "classifier") on a set of test data for which the true values are known.
In the field of machine learning and specifically the problem of statistical classification, a confusion matrix, also known as an error matrix, is a specific table layout that allows visualization of the performance of an algorithm,
typically a supervised learning one (in unsupervised learning it is usually called a matching matrix). Each row of the matrix represents the instances in a predicted class, while each column represents the instances in an actual class (or vice versa).
The name stems from the fact that it makes it easy to see whether the system is confusing two classes (i.e. commonly mislabeling one as another).
It is a special kind of contingency table, with two dimensions ("actual" and "predicted"), and identical sets of "classes" in both dimensions (each combination of dimension and class is a variable in the contingency table).
The Confusion Matrix for the Input Model is As shown:       
''')
        doc.add_picture(self.results_path+'/confusion_matrix.jpeg',width=Inches(5.0))
        doc.add_paragraph('''  ''')
        doc.add_heading("Precision Recall Curve")
        doc.add_paragraph('''
Precision-Recall curves summarize the trade-off between the true positive rate and the positive predictive value for a predictive model using different probability thresholds.
Precision is a ratio of the number of true positives divided by the sum of the true positives and false positives. It describes how good a model is at predicting the positive class. Precision is referred to as the positive predictive value.
A precision-recall curve is a plot of the precision (y-axis) and the recall (x-axis) for different thresholds, much like the ROC curve.
The Precision Recall Curve for the Input Model is a s shown:

        ''')
        doc.add_picture(self.results_path+'/prc.jpeg',width=Inches(5.0))
        doc.save(self.results_path+"/"+file_name+".docx")


