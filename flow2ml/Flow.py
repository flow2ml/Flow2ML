import os
import cv2
import sys ,time
import tensorflowjs as tfjs
import tensorflow as tf
from docx import Document
import matplotlib.pyplot as plt
from .Data_Loader import Data_Loader
from .Filters import Filters
from .Data_Augumentation import Data_Augumentation
from .Image_Quality import Image_Quality

class Flow(Data_Loader,Filters,Data_Augumentation,Image_Quality):
  '''
    This class connects the work flow. It calls various methods from 
    the inherited classes whenever required and completes the workflow.

    Inherits:
      Data_Loader: (Class) Contains various methods to handle data.
      Filters: (Class) Contains methods to apply filters to the given data.
      Data_Augmentation: (Class) Contains various methods to apply augmentation to the given data.
      Image_Quality: (Class) Calculates image quality for each image in the processedDataFolder.
  '''

  def __init__(self,dataset_dir,data_dir):
    '''
      Initializes the main class with all the required information that's 
      to be sent to the methods when they are called.
      Args :
        dataset_dir : (string) the datset containing main folder.
        data_dir : (string) the data containing folder with various labels
                            as sub folders.
    '''
    self.dataset_dir = dataset_dir
    self.data_dir = data_dir
    super().__init__(dataset_dir,data_dir)
    self.classes = self.getClasses()

    try:
      self.classes.remove('.ipynb_checkpoints') 
    except:
      pass
    
    self.results_path = os.path.join(os.getcwd(),"GeneratedReports")
    self.deployment_path = os.path.join(os.getcwd(), "DeployedModels")

    try:
      os.mkdir(self.results_path)
    except:
      raise Exception("Unable to create directory for results.")

    self.num_classes =  len(self.classes)
    self.class_counts = {}
    self.detectBlurred()
    self.categoriesCountPlot()

  def categoriesCountPlot(self):
  
    '''
      Function used to create a countPlot of all categories of images
      Args : None
    '''

    # loop through the folders of each class, store the class and number of images in dictionary
    for folder in self.classes:
      path = os.path.join(self.dataset_dir ,self.data_dir, folder)
      count = len(os.listdir(path))
      self.class_counts[folder] = count
    
    # display the count on the console
    print("Counts per each category")
    for key in self.class_counts.keys():
      print(f"Category {key}: {self.class_counts[key]}")

    # check for class imbalance
    min_class = min(self.class_counts.values())
    max_class = max(self.class_counts.values())
    total_images = sum(self.class_counts.values())
    if max_class - min_class >= 0.5 * total_images:
      create_countplot = input("Class imbalance present. Do you want to continue? (y/n): ")
      if create_countplot.upper() != "Y":
        return

    # create a countplot and temperorily save it as image
    plt.bar(self.class_counts.keys(), self.class_counts.values())
    plt.title('Countplot of categories')
    plt.xlabel('Class')
    plt.ylabel('Number of images')
    plt.savefig(os.path.join(self.results_path, "categoriesCountPlot.png"))

    # create a document in GeneratedReports folder and put the image in it
    try:
      doc = Document()
      doc.add_heading('Categories Countplot')
      doc.add_picture(os.path.join(self.results_path, "categoriesCountPlot.png"))
      doc.save(os.path.join(self.results_path, "categories_countplot.docx"))
    except Exception as e:
      print(f"Unable to create categories countplot document due to {e}")

    # remove the picture file after it has been added to the document
    os.remove(os.path.join(self.results_path, "categoriesCountPlot.png"))

  def detectBlurred(self):
  
    '''
      Function used to detect bluriness in all images provided by the user.
      Args : None
    '''

    blurred = []
    # loop over all images provided by the user
    for folder in self.classes:
      path = os.path.join(self.dataset_dir, self.data_dir, folder)
      # get the specific image name to be read
      for image_name in os.listdir(path):
        image_path = os.path.join(path, image_name)
        # read the image in greyscale mode
        image = cv2.imread(image_path, 0)
        # calculate the focus measure by getting the variance with Laplacian filter
        focus_measure = cv2.Laplacian(image, cv2.CV_64F).var()
        # if the focus is less than a certain threshold, then detect bluriness
        if focus_measure < 100:
          blurred.append(image_path)
    if len(blurred) > 0:
      print(f"Blur detected in the following images: {[os.path.split(i)[1] for i in blurred]}")
      print([os.path.split(i)[1] for i in blurred])
      remove_blurred = input("Do you want to remove those images? (y/n): ")
      if remove_blurred.upper() == "Y":
        for image in blurred:
          os.remove(image)

  def update_progress(self,progress,subStatus):
    
    '''
      Function used to update the progress bar in the console 
      Args :
        progress: float value for representing percentage
    '''
    barLength = 10 # Modify this to change the length of the progress bar
    status=""

    if progress >= 1:
        progress = 1
        status = subStatus +" ...\r\n"

    block = int(round(barLength*progress))
    text = "\rPercent: [{0}] {1}% {2}".format( "#"*block + "-"*(barLength-block), progress*100, status)
    sys.stdout.write(text)
    sys.stdout.flush()

  def applyFilters(self,filters):
    ''' 
      Applies given filters 
      Args : 
        filters : (List) python list containing various filters to 
                         be applied to the image data.
    '''
    self.filters = filters
    # checks to see if the list contains a dictionary with apply_on_augmentation key
    apply_on_augmentation = [i for i in filters if isinstance(i, dict) == True]
    if apply_on_augmentation:
      apply_on_augmentation = apply_on_augmentation[0].get('apply_on_augmentation', False)
      print(apply_on_augmentation)
    
    progress = [(100/len(filters))*i for i in range(0,len(filters)) ]
    progress_i = 0

    for folder in self.classes:
      for filter in filters:
        paths = [os.path.join(self.dataset_dir ,self.data_dir, folder)]
        # if apply_on_augmentation is true, creates a list of folders of augmented images to apply the filters to
        if apply_on_augmentation:
          for operation in apply_on_augmentation:
            paths = [os.path.join(self.dataset_dir ,self.data_dir, folder, operation.title() + "Images")]
        
        # applies filters to all folders in paths (single folder if no apply_on_augmentation)
        for path in paths:
          if filter == "median":
            self.applyMedian(path)

          elif filter == "laplacian":
            self.applylaplacian(path)

          elif filter == "sobelx":
            self.applysobelx(path)

          elif filter == "sobely":
            self.applysobely(path)

          elif filter == "gaussian":
            self.applygaussian(path)
          
          elif filter == "bilateral":
            self.applybilateral(path)
      
        time.sleep(0.1)
        self.update_progress( progress[progress_i]/100.0, f"Filtered all images in {folder}" )
        progress_i += 1

      status1 = f"Applied all filters to {folder} ...\r\n"
      self.update_progress( 100/100.0, f"Filtered all images in {folder}"  )
      progress_i = 0
    print()
    self.visualizeFilters()

  def applyAugmentation(self,operations):
    
    ''' 
      Applies given data augmentation operations 
      Args : 
        operations : (dictonary) python dictionary containing key value pairs of operations and values to be applied to the image data.
    '''
    
    self.operations = operations
    # checks to see if apply_on_filters key if present
    apply_on_filters = self.operations.get('apply_on_filters', False)

    progress = [(100/len(operations))*i for i in range(0,len(operations)) ]
    progress_i = 0

    for folder in self.classes:
      for operation in self.operations:
        paths = [os.path.join(self.dataset_dir ,self.data_dir, folder)]
        # if apply_on_filters is true, creates a list of folders of filtered images to apply the augmentations to
        if apply_on_filters:
          for filter in apply_on_filters:
            paths = [os.path.join(self.dataset_dir ,self.data_dir, folder, filter.title() + "Images")]
        
         # applies augmentations to all folders in paths (single folder if no apply_on_filters)
        for path in paths:
          if operation == "flipped":
            self.applyFlip(path)

          if operation == "rotated":
            self.applyRotate(path)
          
          if operation == "sheared":
            self.applyShear(path)

          if operation == "inverted":
            self.applyInvert(path)

          if operation == "histogramequalised":
            self.applyHistogramEqualization(path)
            
          if operation == "CLAHE":
            self.applyCLAHE(path)
          
          if operation == "cropped":
            self.applyCrop(path)

          if operation == "scaled":
            self.applyScale(path)

          if operation == "zoomed":
            self.applyZoom(path)

          if operation == "greyscale":
            self.applyGreyscale(path)

          if operation == "eroded":
            self.applyErode(path)

          if operation == "dilated":
            self.applyDilate(path)

          if operation == "opened":
            self.applyOpen(path)

          if operation == "closed":
            self.applyClose(path)

          if operation == "thresholded":
            self.applyThreshold(path)

          if operation == "colorspace":
            self.changeColorSpace(path)
            
          if operation == "canny":
            self.applyCanny(path)
          
          if operation == "brightnessenhanced":
            self.applyEnhanceBrightness(path)


        time.sleep(0.1)
        self.update_progress( progress[progress_i]/100.0, f"Augmented all images in {folder}" )
        progress_i += 1

      status1 = f"Applied all operations to {folder} ...\r\n"
      self.update_progress( 100/100.0, f"Augmented all images in {folder}"  )
      progress_i = 0
    
    print()
    self.visualizeAugmentation()

  def getDataset(self,img_dimensions,train_val_split):
    
    '''
      Generates the dataset.
      Moves all the images to a seperate folder and prepares a numpy dataset.

      Args:
        img_dimensions: (tuple) holds dimensions of the image after resizing.
        train_val_split: (float) holds train validation split value.
      
      Returns:
        train_val_dataset: (tuple) contains the numpy ndarrays.
                            (trainData, trainLabels, valData, valLabels).
    '''

    self.update_progress( 0/100.0, "Created Datasets" )  

    # creates processedData folder in self.dataset_dir
    final_data_folder = os.path.join(self.dataset_dir,'processedData')
    try:
      os.mkdir(final_data_folder)    
      print("Creating processedData Folder")
    except Exception as e:
      print(f"Unable to create processedData folder due to {e}")

    # moves all the images to that folder
    for folder in self.classes:
      path = os.path.join(self.dataset_dir ,self.data_dir, folder)
      self.img_label = self.create_dataset(path)
    
    self.update_progress( 50/100.0, "Created Datasets" )

    # Prepare Numpy dataset
    self.dataset = self.prepare_dataset(img_dimensions,train_val_split,self.img_label)

    self.update_progress( 100/100.0,"Created Datasets" ) 

    return self.dataset

  def deployTensorflowModels(self,conversions,model):
    '''Deploy conversion from tensorflow model to tensorflowjs or tflite model'''

    if(conversions['tfjs']==True):
      # Applying the conversion function to the input model and converted tfjs model will be stored in 'trained_models' folder.
      try:
        os.mkdir(os.path.join(self.deployment_path, "tfjs_model"))
      except Exception as e:
        print(f"Failed to create tfjs_model directory due to {e}")
      tfjs.converters.save_keras_model(model, os.path.join(self.deployment_path, "tfjs_model")) 

    if(conversions['tflite']==True):
      TF_LITE_MODEL_FILE_NAME='tf_lite_model.tflite'
      # Defining the convertor
      tf_lite_converter = tf.lite.TFLiteConverter.from_keras_model(model)
      # Applying the convert function
      tflite_model = tf_lite_converter.convert()
      try:
        os.mkdir(os.path.join(self.deployment_path, "tflite_model"))
      except Exception as e:
        print(f"Failed to create tflite_model directory due to {e}")
      open(os.path.join(self.deployment_path, "tflite_model", TF_LITE_MODEL_FILE_NAME),"wb").write(tflite_model)

  def calculateImageQuality(self,image_quality):

    '''
      Function used to calculate image quality for all images in processedData folder
      Args : 
          image_quality : (string) used to calculate quality by BRISQUE or Entropy function
    '''
    processed_data_folder = os.path.join(self.dataset_dir,'processedData')
    self.image_scores = {}
    self.image_quality = image_quality
    self.create_scores_doc(processed_data_folder)