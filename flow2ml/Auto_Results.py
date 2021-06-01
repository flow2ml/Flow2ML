from sklearn.metrics import roc_auc_score, roc_curve
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from sklearn.metrics import plot_confusion_matrix, plot_roc_curve
from sklearn.metrics import plot_precision_recall_curve
from docx import Document
from docx.shared import Inches
import os


class Auto_Results:
    def __init__(self,model,X_test,Y_test):
        '''
        Initialises the class variables.
        '''
        self.model = model          
        # stores the model
        self.x_test = X_test        
        # stores the X_test (test data) 
        self.y_test = Y_test    
        # stores the Y_test (labels)
        try:                        
            #try creating a results directory
            p = os.getcwd()
            p = os.path.join(p,"GeneratedReports")
            self.results_path = p
            if not os.path.exists(self.results_path):
                os.mkdir(self.results_path)
        except Exception as e:
            print(f"Unable to create directory for results due to {e}.")

    def roc_curve(self,figure_name="roc.jpeg"):
        '''
        Plotting the Roc curve for the Input model
        '''
        try:
            plt.figure(figsize=(50, 20))
            plot_roc_curve(self.model, self.x_test, self.y_test) 
            plt.title('ROC Curve')
            plt.savefig(os.path.join(self.results_path, figure_name))
        except Exception as e:
            print(f"Unable to plot roc curve due to {e}")
    
    def confusion_matrix(self,figure_name='confusion_matrix.jpeg'):
        '''
        Plotting the confusion matrix
        '''
        try:
            plt.figure(figsize=(50,20))
            plot_confusion_matrix(self.model,self.x_test,self.y_test)
            plt.title('Confusion Matrix', size = 35)
            plt.savefig(os.path.join(self.results_path, figure_name))
        except Exception as e:
            print(f"Unable to plot confusion matrix due to {e}")
    
    def precision_recall_curve(self,figure_name="prc.jpeg"):
        '''
        Plotting the precision recall curve 
        '''
        try:
            plt.figure(figsize=(50,20))
            plot_precision_recall_curve(self.model,self.x_test,self.y_test)
            plt.title('Precision Recall Curve', size = 35)
            plt.savefig(os.path.join(self.results_path, figure_name))
        except Exception as e:
            print(f"Unable to plot precision recall curve due to {e}")

    def get_results_docx(self,file_name="model_report.docx"):
        '''
        Saves all the plots with their default name and 
        creates a report.docx file in the Results Folder
        '''
        self.roc_curve()
        self.confusion_matrix()
        self.precision_recall_curve()
        try:
            doc = Document()
            doc.add_heading('Model Report')
            doc.add_heading('ROC Curve',level = 2)
            doc.add_paragraph('''
An ROC curve (receiver operating characteristic curve) is a graph showing the performance of a classification model at all classification thresholds.
This curve plots two parameters:
    True Positive Rate
    False Positive Rate
An ROC curve plots TPR vs. FPR at different classification thresholds. Lowering theclassification threshold classifies more items as positive, thus increasing both False Positives and TrPositives.
The ROC curve for the Input Model is As Follows.
            ''')
            doc.add_picture(os.path.join(self.results_path, 'roc.jpeg'),width=Inches(6.0))
            doc.add_paragraph('''  ''')
            doc.add_heading('Confusion Matrix', level=2)
            doc.add_paragraph('''
A confusion matrix is a table that is often used to describe the performance of a classification model (or "classifier") on a set of test data for which the true values are known.
In the field of machine learning and specifically the problem of statistical classification, a confusion matrix, also known as an error matrix, is a specific table layout that allows visualization of the performance of an algorithm, typically a supervised learning one (in unsupervised learning it is usually called a matching matrix). Each row of the matrix represents the instances in a predicted class, while each column represents the instances in an actual class (or vice versa).
The name stems from the fact that it makes it easy to see whether the system is confusing two classes (i.e. commonly mislabeling one as another).
It is a special kind of contingency table, with two dimensions ("actual" and "predicted"), and identical sets of "classes" in both dimensions (each combination of dimension and class is a variable in the contingency table).
The Confusion Matrix for the Input Model is As shown:       
            ''')
            doc.add_picture(os.path.join(self.results_path, 'confusion_matrix.jpeg'),width=Inches(5.0))
            doc.add_paragraph('''  ''')
            doc.add_heading("Precision Recall Curve")
            doc.add_paragraph('''
Precision-Recall curves summarize the trade-off between the true positive rate and the positive predictive value for a predictive model using different probability thresholds.
Precision is a ratio of the number of true positives divided by the sum of the true positives and false positives. It describes how good a model is at predicting the positive claPrecision is referred to as the positive predictive value.
A precision-recall curve is a plot of the precision (y-axis) and the recall (x-axis) for different thresholds, much like the ROC curve.
The Precision Recall Curve for the Input Model is a s shown:
            ''')
            doc.add_picture(os.path.join(self.results_path, 'prc.jpeg'),width=Inches(5.0))
            doc.save(os.path.join(self.results_path, file_name))
            # remove the saved pictures after they are added to the document
            os.remove(os.path.join(self.results_path, 'roc.jpeg'))
            os.remove(os.path.join(self.results_path, 'confusion_matrix.jpeg'))
            os.remove(os.path.join(self.results_path, 'prc.jpeg'))
        except Exception as e:
            print(f"Unable to create results document due to {e}")