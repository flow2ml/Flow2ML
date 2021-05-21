from sklearn.metrics import roc_curve
import matplotlib.pyplot as plt
from sklearn.metrics import roc_curve,auc
from sklearn.metrics import precision_recall_curve
from docx import Document
from docx.shared import Inches
import os
import tensorflow as tf
import seaborn as sn

class Tf_Results:
    def __init__(self,model,validtion_generator):
        '''
        Initialises the class variables.
        '''
        self.model = model          #stores the model
        self.validation_generator = validtion_generator  #stores the validation generator object   
        try:                        #try creating a Results directory and show excpetion if it already exists 
            p = os.getcwd()
            p = os.path.join(p,"Results")
            self.results_path = p
            os.mkdir(self.results_path)
        except:
            raise Exception("Unable to create directory for results.")

    def tf_roc_curve(self,figure_name="roc.jpeg",labels=None):
        '''
        Plotting the Roc curve for the Input model
        '''
        if labels == None:
            labels = self.validation_generator.classes
        preds = self.model.predict(self.validation_generator,verbose=1)
        fpr, tpr, _ = roc_curve(labels, preds)
        roc_auc = auc(fpr, tpr)
        lw = 2
        plt.figure(figsize=(50, 20))
        plt.plot(fpr, tpr, color='darkorange',
        lw=lw, label='ROC curve (area = %0.2f)' % roc_auc)
        try:
            plt.plot([0, 1], [0, 1], color='navy', lw=lw, linestyle='--')
            plt.xlabel('False Positive Rate')
            plt.ylabel('True Positive Rate')
            plt.title('Receiver operating characteristic')
            plt.legend(loc="lower right")
            plt.savefig(self.results_path+'/'+figure_name)
        except Exception as e:
            print("Unable to plot roc curve due to {e}")
    
    def tf_confusion_matrix(self):
        '''
        Plotting the confusion matrix
        '''
        preds = self.model.predict(self.validation_generator,verbose=1)
        labels = self.validation_generator.classes
        df_cm = tf.math.confusion_matrix(
            labels, preds, num_classes=None, weights=None, dtype=tf.dtypes.int32,
            name=None)
        try:
            plt.figure(figsize=(50,20))
            sn.heatmap(df_cm, annot=True, annot_kws={"size": 35})
            plt.title('Confusion Matrix', size = 35)
            plt.savefig(self.results_path+'/confusion_matrix.jpeg')
        except Exception as e:
            print("Unable to plot confusion matrix due to {e}")
    
    def tf_precision_recall_curve(self,figure_name="prc.jpeg",labels=None):
        '''
        Plotting the precision recall curve 
        '''
        if labels == None:
            labels = self.validation_generator.classes
        preds = self.model.predict(self.validation_generator,verbose=1)
        pr, tpr, _ = precision_recall_curve(labels, preds)
        lw = 2
        try:
            plt.figure(figsize=(50,20))
            plt.plot(tpr, pr, color='darkorange',
            lw=lw)
            plt.xlabel('Recall')
            plt.ylabel('Precision')
            plt.title('Precision Recall Curve')
            plt.savefig(self.results_path+'/'+figure_name)
        except Exception as e:
            print("Unable to plot precision recall curve due to {e}")

    def tf_get_results_docx(self,file_name="report.docx"):
        '''
        Saves all the plots with their default name and 
        creates a report.docx file in the Results Folder
        '''
        self.tf_roc_curve()
        self.tf_confusion_matrix()
        self.tf_precision_recall_curve()
        try:
            doc = Document()
            doc.add_heading('Model Report')
            doc.add_heading('ROC Curve',level = 2)
            doc.add_paragraph('''
                An ROC curve (receiver operating characteristic curve) is a graph showing the performance of a classification model at all classification thresholds.
                This curve plots two parameters:
                        True Positive Rate
                        False Positive Rate
                An ROC curve plots TPR vs. FPR at different classification thresholds. Lowering the classification threshold classifies more items as positive, thus increasing both False Positives and True Positives.
                The ROC curve for the Input Model is As Follows.
                ''')
            doc.add_picture(self.results_path+'/roc.jpeg',width=Inches(6.0))
            doc.add_paragraph('''  ''')
            doc.add_heading('Confusion Matrix', level=2)
            doc.add_paragraph('''
                A confusion matrix is a table that is often used to describe the performance of a classification model (or "classifier") on a set of test data for which the true values are known.
                In the field of machine learning and specifically the problem of statistical classification, a confusion matrix, also known as an error matrix, is a specific table layout that allows visualization of the performance of an algorithm,
                typically a supervised learning one (in unsupervised learning it is usually called a matching matrix). Each row of the matrix represents the instances in a predicted class, while each column represents the instances in an actual class (or vice versa).
                The name stems from the fact that it makes it easy to see whether the system is confusing two classes (i.e. commonly mislabeling one as another).
                It is a special kind of contingency table, with two dimensions ("actual" and "predicted"), and identical sets of "classes" in both dimensions (each combination of dimension and class is a variable in the contingency table).
                The Confusion Matrix for the Input Model is As shown:       
                ''')
            doc.add_picture(self.results_path+'/confusion_matrix.jpeg',width=Inches(5.0))
            doc.add_paragraph('''  ''')
            doc.add_heading("Precision Recall Curve")
            doc.add_paragraph('''
                Precision-Recall curves summarize the trade-off between the true positive rate and the positive predictive value for a predictive model using different probability thresholds.
                Precision is a ratio of the number of true positives divided by the sum of the true positives and false positives. It describes how good a model is at predicting the positive class. Precision is referred to as the positive predictive value.
                A precision-recall curve is a plot of the precision (y-axis) and the recall (x-axis) for different thresholds, much like the ROC curve.
                The Precision Recall Curve for the Input Model is a s shown:

                        ''')
            doc.add_picture(self.results_path+'/prc.jpeg',width=Inches(5.0))
            doc.save(self.results_path+"/"+file_name)
        except Exception as e:
            print("Unable to create results document due to {e}")