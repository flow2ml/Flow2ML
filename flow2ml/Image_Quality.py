import imquality.brisque as brisque
import cv2
from docx import Document
import os
import numpy as np

class Image_Quality:
	'''
		Class to calculate image quality score for all images that are stored in ProcessedData folder
	'''

	def __init__(self, image_quality):
		'''
		Initializes various attributes regarding to the object.
		Args : 
			image_quality : (string) used to calculate quality by BRISQUE or Entropy function
		'''
		self.image_scores = {}
		self.image_quality = image_quality

	def generate_img_scores(self, path):
		''' 
		Generates the image quality scores for all images in the given folder using BRISQUE and stores it in a dictionary.
		Args : 
			path : (string) path to the processedData folder.
		'''
		
		for i in os.listdir(path):
			img_path = os.path.join(path, i)
			img = cv2.imread(img_path, 0)
			
			if self.image_quality == "brisque":
				try:
					img_score = brisque.score(img)
				except Exception as e:
					print(f"Unable to calculate BRISQUE socre due to {e}")
			
			elif self.image_quality == "entropy":
				try:
					entropy = []
					hist = cv2.calcHist([img], [0], None, [256], [0, 255])
					total_pixel = img.shape[0] * img.shape[1]
					for item in hist:
						probability = item / total_pixel
						if probability == 0:
							en = 0
						else:
							en = -1 * probability * (np.log(probability) / np.log(2))
						entropy.append(en)
					img_score = np.sum(entropy)[0]
				except Exception as e:
					print(f"Unable to calculate Entropy socre due to {e}")
			
			# add score and image name to dictionary after rounding it off
			self.image_scores[i] = format(img_score, '.2f')

	def create_scores_doc(self, path):
		''' 
		Calculates image scores for the processedData folder and creates a document containing this information in GeneratedReports
		Args : 
			path : (string) path to the processedData folder.
		'''
		self.generate_img_scores(path)
		try:
			doc = Document()
			doc.add_heading('Image Quality Report')
			# loop over all images in dictionary and add their names and scores to the document.
			for image in self.image_scores:
				doc.add_paragraph(f"Image {image}: Quality Score {self.image_scores[image]}")
			doc.save(os.path.join(self.results_path, "image_quality_report.docx"))
		except Exception as e:
			print(f"Unable to create image quality document due to {e}")